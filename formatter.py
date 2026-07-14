"""
Carry-On Confidence — DOCX formatter.

Takes the content dict produced by generator.generate_worksheet() and renders
it into two DOCX files: the student worksheet and (if applicable) an answer key.
Pure presentation layer — no LLM calls, no data fetching.
"""

import os
import re
import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

_KEYWORD_SEPARATOR = "______________________________________________________________________________________________"


def _add_paragraph(doc, text="", bold=False, italic=False, size=12, alignment=None):
    """Shared helper. Adds a paragraph with a single run. Returns the paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if alignment is not None:
        para.alignment = alignment
    return para


def _render_keyword_page(doc, page_content):
    """Parses 3 keyword blocks from page_content and renders them into doc."""
    blocks = page_content.split(_KEYWORD_SEPARATOR)
    blocks = [b.strip() for b in blocks if b.strip()]

    for block in blocks[:3]:
        lines = [l for l in block.splitlines() if l.strip()]

        for line in lines:
            stripped = line.strip()

            if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
                # Word + POS line: **word (pos)**
                text = stripped[2:-2]
                _add_paragraph(doc, text, bold=True, size=12)

            elif stripped.startswith("**English:**"):
                definition = stripped[len("**English:**"):].strip()
                para = doc.add_paragraph()
                bold_run = para.add_run("English: ")
                bold_run.bold = True
                bold_run.font.size = Pt(12)
                plain_run = para.add_run(definition)
                plain_run.font.size = Pt(12)

            elif stripped.startswith("**Korean:**"):
                definition = stripped[len("**Korean:**"):].strip()
                para = doc.add_paragraph()
                bold_run = para.add_run("Korean: ")
                bold_run.bold = True
                bold_run.font.size = Pt(12)
                plain_run = para.add_run(definition)
                plain_run.font.size = Pt(12)

            elif stripped.startswith("**Synonyms:**"):
                synonyms = stripped[len("**Synonyms:**"):].strip()
                para = doc.add_paragraph()
                bold_run = para.add_run("Synonyms: ")
                bold_run.bold = True
                bold_run.font.size = Pt(12)
                plain_run = para.add_run(synonyms)
                plain_run.font.size = Pt(12)

            elif stripped.startswith("**My sentence:**"):
                para = doc.add_paragraph()
                bold_run = para.add_run("My sentence: ")
                bold_run.bold = True
                bold_run.font.size = Pt(12)

        # Two writing blank lines
        _add_paragraph(doc, "_" * 100)
        _add_paragraph(doc, "_" * 100)

        # Spacing between blocks
        _add_paragraph(doc)

    doc.add_page_break()


def _render_role_play_page(doc, page_content):
    """Renders one role play card into doc. Adds a page break at the end."""
    for line in page_content.splitlines():
        if not line.strip():
            doc.add_paragraph()
            continue

        if "**" in line:
            parts = re.split(r'(\*\*.+?\*\*)', line)
            para = doc.add_paragraph()
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(12)
                else:
                    run = para.add_run(part)
                    run.font.size = Pt(12)
        else:
            _add_paragraph(doc, line, size=12)

    doc.add_page_break()


def format_worksheet(content: dict, level: int, level_config: dict, topic: str, location: str, _doc=None) -> str:
    """Public entry point. Renders Page 1 of the worksheet into a Document."""
    doc = _doc if _doc is not None else Document()

    # 2. Title image
    title_img = os.path.join(os.path.dirname(__file__), "assets", "carry-on-title.png")
    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(title_img, width=Inches(6.78))
    except Exception:
        _add_paragraph(doc, "Carry-On Confidence", bold=True, size=24, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _add_paragraph(doc, "Fake It Till You Make the Flight", italic=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 3. Level line
    para = doc.add_paragraph()
    bold_run = para.add_run("Level: ")
    bold_run.bold = True
    bold_run.font.size = Pt(12)
    plain_run = para.add_run(level_config["cefr"] + " — " + level_config["label"])
    plain_run.font.size = Pt(12)

    # 4. Name field
    para = doc.add_paragraph()
    bold_run = para.add_run("Name: ")
    bold_run.bold = True
    bold_run.font.size = Pt(12)
    plain_run = para.add_run("_" * 40)
    plain_run.font.size = Pt(12)

    # 5. Date field
    today = datetime.datetime.now().strftime("%B %d, %Y").replace(" 0", " ")
    para = doc.add_paragraph()
    bold_run = para.add_run("Date: ")
    bold_run.bold = True
    bold_run.font.size = Pt(12)
    plain_run = para.add_run(today)
    plain_run.font.size = Pt(12)

    # 6. Section heading
    _add_paragraph(doc, "Let's Talk!", bold=True, size=14)

    # 7. Small talk questions
    for line in content["page1"].splitlines():
        stripped = line.strip()
        if stripped:
            _add_paragraph(doc, stripped, size=12)
            _add_paragraph(doc, "_" * 100, size=12)

    # 8. Page break
    doc.add_page_break()

    # PAGE 2 — Keywords 1-3
    _add_paragraph(doc, "Keywords", bold=True, size=14)
    _render_keyword_page(doc, content["page2"])

    # PAGE 3 — Keywords 4-6
    _add_paragraph(doc, "Keywords", bold=True, size=14)
    _render_keyword_page(doc, content["page3"])

    # PAGE 4 — Role Play Card 1
    _add_paragraph(doc, "Role Play", bold=True, size=14)
    _render_role_play_page(doc, content["page4"])

    # PAGE 5 — Role Play Card 2
    _add_paragraph(doc, "Role Play", bold=True, size=14)
    _render_role_play_page(doc, content["page5"])

    # PAGES 6-7 — Exercises (natural flow, no forced page break between)
    _add_paragraph(doc, "Exercises", bold=True, size=14)
    exercise_num = 0
    for line in content["page6"].splitlines():
        if line.strip() == "===EXERCISES_END===":
            break
        if line.strip() == "---":
            continue
        if not line.strip():
            doc.add_paragraph()
            continue
        if re.match(r'^coc_\d+:', line.strip()):
            exercise_num += 1
            label = re.sub(r'^coc_\d+:\s*', '', line.strip())
            _add_paragraph(doc, str(exercise_num) + ". " + label, bold=True, size=12)
        elif line.strip().startswith("Instructions:"):
            _add_paragraph(doc, line.strip(), italic=True, size=12)
        elif line.strip().startswith("___"):
            _add_paragraph(doc, line.strip(), size=12)
        else:
            _add_paragraph(doc, line.strip(), size=12)

    # PAGE 8 — Homework
    doc.add_page_break()
    _add_paragraph(doc, "Homework", bold=True, size=14)
    for line in content["page8"].splitlines():
        if line.strip():
            _add_paragraph(doc, line.strip(), size=12)
    for _ in range(17):
        _add_paragraph(doc, "_" * 100)

    # Filename and save
    filename = "coc_" + location + "_" + topic + "_L" + str(level) + "_" + datetime.datetime.now().strftime("%Y%m%d") + ".docx"
    if _doc is not None:
        return "test_mode"
    os.makedirs("output", exist_ok=True)
    out_path = os.path.join("output", filename)
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    import sys

    MOCK_PAGE = """**boarding pass (noun)**
**English:** The document you need to get on the plane.
**Korean:** 비행기 탑승을 위해 필요한 서류입니다.
**Synonyms:** ticket · pass · travel document
**My sentence:** ______________________________

______________________________________________________________________________________________
**luggage (noun)**
**English:** The bags and suitcases you travel with.
**Korean:** 여행할 때 가지고 다니는 가방과 짐입니다.
**Synonyms:** baggage · suitcase · bags
**My sentence:** ______________________________

______________________________________________________________________________________________
**itinerary (noun)**
**English:** A detailed plan of your trip including dates and destinations.
**Korean:** 날짜와 목적지가 포함된 여행 계획서입니다.
**Synonyms:** schedule · plan · route
**My sentence:** ______________________________

______________________________________________________________________________________________"""

    # Test 1 — _add_paragraph()
    try:
        doc1 = Document()
        para = _add_paragraph(doc1, "Test bold paragraph", bold=True, size=14)
        assert para.runs[0].text == "Test bold paragraph", "Text mismatch"
        assert para.runs[0].bold is True, "Expected bold"
        print("PASS: _add_paragraph")
    except Exception as e:
        print("FAIL: _add_paragraph —", e)
        sys.exit(1)

    # Test 2 — _render_keyword_page()
    try:
        doc2 = Document()
        _render_keyword_page(doc2, MOCK_PAGE)
        all_text = " ".join(p.text for p in doc2.paragraphs)
        assert "boarding pass" in all_text, "Expected 'boarding pass' in paragraphs"
        assert "비행기 탑승" in all_text, "Expected Korean text in paragraphs"
        print("PASS: _render_keyword_page")
    except Exception as e:
        print("FAIL: _render_keyword_page —", e)
        sys.exit(1)

    MOCK_KEYWORDS = """**boarding pass (noun)**
**English:** The document you need to get on the plane.
**Korean:** 비행기 탑승을 위해 필요한 서류입니다.
**Synonyms:** ticket · pass · travel document
**My sentence:** ______________________________

______________________________________________________________________________________________
**luggage (noun)**
**English:** The bags and suitcases you travel with.
**Korean:** 여행할 때 가지고 다니는 가방과 짐입니다.
**Synonyms:** baggage · suitcase · bags
**My sentence:** ______________________________

______________________________________________________________________________________________
**itinerary (noun)**
**English:** A detailed plan of your trip including dates and destinations.
**Korean:** 날짜와 목적지가 포함된 여행 계획서입니다.
**Synonyms:** schedule · plan · route
**My sentence:** ______________________________

______________________________________________________________________________________________"""

    MOCK_EXERCISES_STUB = "coc_001: Warm Up\nInstructions: Answer the question.\nWhat do you do at the airport?\n_______________\n===EXERCISES_END==="
    MOCK_HOMEWORK_STUB = "Write about your last trip."

    # Test 3 — format_worksheet() Page 1
    try:
        mock_content = {
            "page1": "What is your favourite travel memory?\nWhere do you want to travel next?\nHave you ever had a travel disaster?",
            "page2": MOCK_KEYWORDS,
            "page3": MOCK_KEYWORDS,
            "page4": MOCK_KEYWORDS,
            "page5": MOCK_KEYWORDS,
            "page6": MOCK_EXERCISES_STUB,
            "page8": MOCK_HOMEWORK_STUB,
        }
        mock_level_config = {"cefr": "B2", "label": "Upper Intermediate"}
        test_doc = Document()
        format_worksheet(mock_content, 6, mock_level_config, "airport", "tokyo", _doc=test_doc)
        all_text = " ".join(p.text for p in test_doc.paragraphs)
        assert "Let's Talk!" in all_text, "Expected 'Let's Talk!' in paragraphs"
        assert "What is your favourite travel memory?" in all_text, "Expected question in paragraphs"
        assert "B2" in all_text, "Expected 'B2' in paragraphs"
        print("PASS: format_worksheet Page 1")
    except Exception as e:
        print("FAIL: format_worksheet Page 1 —", e)
        sys.exit(1)

    # Test 4 — format_worksheet() Pages 2-3
    try:
        mock_content_4 = {
            "page1": "Have you ever missed a flight?",
            "page2": MOCK_KEYWORDS,
            "page3": MOCK_KEYWORDS,
            "page4": MOCK_KEYWORDS,
            "page5": MOCK_KEYWORDS,
            "page6": MOCK_EXERCISES_STUB,
            "page8": MOCK_HOMEWORK_STUB,
        }
        test_doc4 = Document()
        format_worksheet(mock_content_4, 6, {"cefr": "B2", "label": "Upper Intermediate"}, "airport", "tokyo", _doc=test_doc4)
        para_texts = [p.text for p in test_doc4.paragraphs]
        assert para_texts.count("Keywords") >= 2, "Expected 'Keywords' at least twice"
        assert any("비행기 탑승" in t for t in para_texts), "Expected '비행기 탑승' in paragraphs"
        assert any("날짜와 목적지" in t for t in para_texts), "Expected '날짜와 목적지' in paragraphs"
        print("PASS: format_worksheet Pages 2-3")
    except Exception as e:
        print("FAIL: format_worksheet Pages 2-3 —", e)
        sys.exit(1)

    MOCK_ROLE_PLAY = """🎭 SITUATION: Lost at the Airport

📍 **Location:** Arrival hall, Incheon International Airport
👤 **You are:** A Korean traveler who just landed and cannot find the exit
👥 **Your partner is:** An airport information desk worker

The situation:
You landed 30 minutes ago but cannot find the immigration queue. You are confused
and a little panicked. You need to ask for help politely and clearly.

**Your goal:**
Get clear directions to immigration and thank the worker.

Try to use these words:
boarding pass · itinerary · luggage"""

    # Test 5 — format_worksheet() Pages 4-5
    try:
        mock_content_5 = {
            "page1": "Have you ever missed a flight?",
            "page2": MOCK_KEYWORDS,
            "page3": MOCK_KEYWORDS,
            "page4": MOCK_ROLE_PLAY,
            "page5": MOCK_ROLE_PLAY,
            "page6": MOCK_EXERCISES_STUB,
            "page8": MOCK_HOMEWORK_STUB,
        }
        test_doc5 = Document()
        format_worksheet(mock_content_5, 6, {"cefr": "B2", "label": "Upper Intermediate"}, "airport", "tokyo", _doc=test_doc5)
        para_texts = [p.text for p in test_doc5.paragraphs]
        assert any("🎭 SITUATION:" in t for t in para_texts), "Expected '🎭 SITUATION:' in paragraphs"
        assert any("Role Play" in t for t in para_texts), "Expected 'Role Play' in paragraphs"
        assert para_texts.count("Role Play") >= 2, "Expected 'Role Play' at least twice"
        print("PASS: format_worksheet Pages 4-5")
    except Exception as e:
        print("FAIL: format_worksheet Pages 4-5 —", e)
        sys.exit(1)

    MOCK_EXERCISES = """coc_001: What Do You Say When...?
Instructions: Read the situation and write what you would say.
You are at the check-in counter and the agent asks for your passport.
What do you say?
_______________________________________________

---
coc_004: Fill the Silence
Instructions: The conversation has stalled. Write what you say next.
The agent has gone quiet after scanning your boarding pass.
_______________________________________________

===EXERCISES_END===
This line must not appear in the document."""

    MOCK_HOMEWORK = "Write a short paragraph about a time you had to ask for help while travelling."

    # Test 6 — format_worksheet() Pages 6-8
    try:
        mock_content_6 = {
            "page1": "Have you ever missed a flight?",
            "page2": MOCK_KEYWORDS,
            "page3": MOCK_KEYWORDS,
            "page4": MOCK_ROLE_PLAY,
            "page5": MOCK_ROLE_PLAY,
            "page6": MOCK_EXERCISES,
            "page8": MOCK_HOMEWORK,
        }
        test_doc6 = Document()
        result6 = format_worksheet(mock_content_6, 6, {"cefr": "B2", "label": "Upper Intermediate"}, "airport", "tokyo", _doc=test_doc6)
        assert result6 == "test_mode", "Expected return value 'test_mode'"
        para_texts = [p.text for p in test_doc6.paragraphs]
        assert any("Exercises" in t for t in para_texts), "Expected 'Exercises' in paragraphs"
        assert any("Homework" in t for t in para_texts), "Expected 'Homework' in paragraphs"
        assert not any("===EXERCISES_END===" in t for t in para_texts), "'===EXERCISES_END===' must not appear"
        assert not any("This line must not appear in the document" in t for t in para_texts), "Sentinel line must not appear"
        assert any("What Do You Say When" in t for t in para_texts), "Expected exercise title in paragraphs"
        blank_count = sum(1 for t in para_texts if t == "_" * 100)
        assert blank_count >= 17, f"Expected at least 17 blank lines, got {blank_count}"
        print("PASS: format_worksheet Pages 6-8")
    except Exception as e:
        print("FAIL: format_worksheet Pages 6-8 —", e)
        sys.exit(1)

    # Test 7 — format_worksheet() save to disk
    try:
        mock_content_7 = {
            "page1": "Have you ever missed a flight?",
            "page2": MOCK_KEYWORDS,
            "page3": MOCK_KEYWORDS,
            "page4": MOCK_ROLE_PLAY,
            "page5": MOCK_ROLE_PLAY,
            "page6": MOCK_EXERCISES,
            "page8": MOCK_HOMEWORK,
        }
        mock_level_config_7 = {"cefr": "B2", "label": "Upper Intermediate"}
        result7 = format_worksheet(mock_content_7, 6, mock_level_config_7, "tokyo", "airport")
        assert isinstance(result7, str) and result7.endswith(".docx"), f"Expected .docx path, got {result7!r}"
        assert os.path.exists(result7), f"Expected file to exist: {result7}"
        print("PASS: format_worksheet save to disk")
    except Exception as e:
        print("FAIL: format_worksheet save to disk —", e)
        sys.exit(1)
